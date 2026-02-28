"""
Document Processing Utilities
Handles text extraction from various document formats.
"""

import streamlit as st


def extract_text_from_pdf(file) -> str:
    """Extract text from a PDF file."""
    try:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        return "\n\n".join(text_parts)

    except ImportError:
        # Fallback to PyPDF2
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)

        except Exception as e:
            st.error(f"PDF extraction failed: {str(e)}")
            return ""


def extract_text_from_docx(file) -> str:
    """Extract text from a Word document."""
    try:
        from docx import Document

        doc = Document(file)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    except Exception as e:
        st.error(f"DOCX extraction failed: {str(e)}")
        return ""


def extract_text_from_txt(file) -> str:
    """Extract text from a plain text file."""
    try:
        # Try different encodings
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                file.seek(0)
                content = file.read()
                if isinstance(content, bytes):
                    return content.decode(encoding)
                return content
            except UnicodeDecodeError:
                continue

        return ""

    except Exception as e:
        st.error(f"Text extraction failed: {str(e)}")
        return ""


def extract_text_from_file(uploaded_file) -> str | None:
    """
    Extract text from an uploaded file based on its type.

    Args:
        uploaded_file: Streamlit UploadedFile object

    Returns:
        Extracted text as string, or None if extraction fails
    """
    if uploaded_file is None:
        return None

    file_type = uploaded_file.type
    file_name = uploaded_file.name.lower()

    # Reset file pointer
    uploaded_file.seek(0)

    # Determine file type and extract accordingly
    if file_type == "application/pdf" or file_name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)

    elif (
        file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or file_name.endswith(".docx")
    ):
        return extract_text_from_docx(uploaded_file)

    elif file_type == "text/plain" or file_name.endswith(".txt"):
        return extract_text_from_txt(uploaded_file)

    else:
        st.error(f"Unsupported file type: {file_type}")
        return None


def chunk_text(text: str, chunk_size: int = 4000, overlap: int = 200) -> list[str]:
    """
    Split text into overlapping chunks for processing large documents.

    Args:
        text: The text to chunk
        chunk_size: Maximum characters per chunk
        overlap: Number of overlapping characters between chunks

    Returns:
        List of text chunks
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # Try to break at a sentence or paragraph boundary
        if end < len(text):
            # Look for paragraph break
            para_break = text.rfind("\n\n", start, end)
            if para_break > start + chunk_size // 2:
                end = para_break + 2
            else:
                # Look for sentence break
                for punct in [". ", ".\n", "? ", "?\n", "! ", "!\n"]:
                    sent_break = text.rfind(punct, start, end)
                    if sent_break > start + chunk_size // 2:
                        end = sent_break + len(punct)
                        break

        chunks.append(text[start:end].strip())
        start = end - overlap

    return chunks


def estimate_tokens(text: str) -> int:
    """
    Estimate the number of tokens in a text.
    Rough estimate: ~4 characters per token for English text.
    """
    return len(text) // 4
