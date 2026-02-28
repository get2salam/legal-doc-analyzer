"""Document parsers for various file formats.

Supports PDF, DOCX, TXT, and HTML documents. Each parser extracts
structured text with page number annotations for clause location tracking.
"""

from __future__ import annotations

import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class ParsedPage:
    """A single page of parsed content."""

    page_number: int
    text: str
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Structured output from document parsing."""

    filename: str
    pages: list[ParsedPage] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        """Concatenate all pages into a single text string."""
        return "\n\n".join(page.text for page in self.pages)

    @property
    def page_count(self) -> int:
        """Total number of pages."""
        return len(self.pages)

    def get_page_for_char_offset(self, offset: int) -> int | None:
        """Given a character offset in full_text, return the page number."""
        current = 0
        for page in self.pages:
            page_end = current + len(page.text) + 2  # +2 for \n\n separator
            if offset < page_end:
                return page.page_number
            current = page_end
        return self.pages[-1].page_number if self.pages else None


class DocumentParser(ABC):
    """Abstract base class for document parsers.

    All parsers must implement the ``parse`` method which takes a file path
    and returns a ``ParsedDocument`` with structured, page-annotated text.
    """

    supported_extensions: tuple[str, ...] = ()

    def can_handle(self, path: Path) -> bool:
        """Check if this parser can handle the given file."""
        return path.suffix.lower() in self.supported_extensions

    @abstractmethod
    def parse(self, path: Path) -> ParsedDocument:
        """Parse a document file and return structured text with page info.

        Args:
            path: Path to the document file.

        Returns:
            ParsedDocument with pages and metadata.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file format is unsupported or corrupted.
        """
        ...

    def _validate_path(self, path: Path) -> None:
        """Validate that the file exists and has a supported extension."""
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        if not self.can_handle(path):
            raise ValueError(
                f"Unsupported file extension '{path.suffix}' for {self.__class__.__name__}. "
                f"Supported: {self.supported_extensions}"
            )


class TextParser(DocumentParser):
    """Parser for plain text files.

    Treats the entire file as a single page, or splits on form-feed
    characters (``\\f``) if present.
    """

    supported_extensions = (".txt", ".text", ".md")

    def parse(self, path: Path) -> ParsedDocument:
        """Parse a plain text file.

        Args:
            path: Path to the text file.

        Returns:
            ParsedDocument with text content.
        """
        self._validate_path(path)
        text = path.read_text(encoding="utf-8", errors="replace")

        # Split on form-feed characters if present
        raw_pages = text.split("\f") if "\f" in text else [text]

        pages = [
            ParsedPage(page_number=i + 1, text=page_text.strip())
            for i, page_text in enumerate(raw_pages)
            if page_text.strip()
        ]

        if not pages:
            pages = [ParsedPage(page_number=1, text="")]

        return ParsedDocument(
            filename=path.name,
            pages=pages,
            metadata={"format": "text", "encoding": "utf-8"},
        )


class PDFParser(DocumentParser):
    """Parser for PDF documents using pdfplumber.

    Extracts text page-by-page with accurate layout preservation.
    Falls back gracefully if pdfplumber is not installed.
    """

    supported_extensions = (".pdf",)

    def parse(self, path: Path) -> ParsedDocument:
        """Parse a PDF file using pdfplumber.

        Args:
            path: Path to the PDF file.

        Returns:
            ParsedDocument with per-page text.

        Raises:
            ImportError: If pdfplumber is not installed.
        """
        self._validate_path(path)

        try:
            import pdfplumber
        except ImportError as exc:
            raise ImportError(
                "pdfplumber is required for PDF parsing. Install it with: pip install pdfplumber"
            ) from exc

        pages = []
        metadata: dict = {"format": "pdf"}

        with pdfplumber.open(str(path)) as pdf:
            metadata["page_count"] = len(pdf.pages)
            metadata["pdf_metadata"] = pdf.metadata or {}

            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                pages.append(
                    ParsedPage(
                        page_number=i + 1,
                        text=text.strip(),
                        metadata={
                            "width": page.width,
                            "height": page.height,
                        },
                    )
                )

        return ParsedDocument(filename=path.name, pages=pages, metadata=metadata)


class DOCXParser(DocumentParser):
    """Parser for Microsoft Word DOCX files using python-docx.

    Extracts paragraph text and maps paragraphs to approximate page
    positions (DOCX does not have native page boundaries in the XML).
    """

    supported_extensions = (".docx",)

    # Rough estimate: paragraphs per page for page-number approximation
    _PARAGRAPHS_PER_PAGE = 25

    def parse(self, path: Path) -> ParsedDocument:
        """Parse a DOCX file using python-docx.

        Args:
            path: Path to the DOCX file.

        Returns:
            ParsedDocument with text grouped into approximate pages.

        Raises:
            ImportError: If python-docx is not installed.
        """
        self._validate_path(path)

        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX parsing. Install it with: pip install python-docx"
            ) from exc

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Group paragraphs into approximate pages
        pages = []
        for i in range(0, max(len(paragraphs), 1), self._PARAGRAPHS_PER_PAGE):
            chunk = paragraphs[i : i + self._PARAGRAPHS_PER_PAGE]
            page_text = "\n\n".join(chunk)
            if page_text.strip():
                pages.append(
                    ParsedPage(
                        page_number=(i // self._PARAGRAPHS_PER_PAGE) + 1,
                        text=page_text,
                    )
                )

        if not pages:
            pages = [ParsedPage(page_number=1, text="")]

        # Extract core properties if available
        metadata: dict = {"format": "docx", "paragraph_count": len(paragraphs)}
        try:
            props = doc.core_properties
            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
        except Exception:
            pass

        return ParsedDocument(filename=path.name, pages=pages, metadata=metadata)


class HTMLParser(DocumentParser):
    """Parser for HTML documents.

    Strips HTML tags and extracts readable text content. Uses Python's
    built-in html.parser module â€” no external dependencies required.
    """

    supported_extensions = (".html", ".htm")

    def parse(self, path: Path) -> ParsedDocument:
        """Parse an HTML file by stripping tags and extracting text.

        Args:
            path: Path to the HTML file.

        Returns:
            ParsedDocument with cleaned text content.
        """
        self._validate_path(path)

        raw_html = path.read_text(encoding="utf-8", errors="replace")
        text = self._strip_html(raw_html)

        pages = [ParsedPage(page_number=1, text=text.strip())]

        return ParsedDocument(
            filename=path.name,
            pages=pages,
            metadata={"format": "html"},
        )

    @staticmethod
    def _strip_html(html: str) -> str:
        """Remove HTML tags and decode entities to produce plain text."""
        import html as html_module
        from html.parser import HTMLParser as StdHTMLParser

        class _TextExtractor(StdHTMLParser):
            def __init__(self) -> None:
                super().__init__()
                self.parts: list[str] = []
                self._skip = False

            def handle_starttag(self, tag: str, attrs: list) -> None:
                if tag in ("script", "style", "head"):
                    self._skip = True

            def handle_endtag(self, tag: str) -> None:
                if tag in ("script", "style", "head"):
                    self._skip = False
                if tag in ("p", "div", "br", "h1", "h2", "h3", "h4", "h5", "h6", "li", "tr"):
                    self.parts.append("\n")

            def handle_data(self, data: str) -> None:
                if not self._skip:
                    self.parts.append(data)

        extractor = _TextExtractor()
        extractor.feed(html)
        text = "".join(extractor.parts)
        text = html_module.unescape(text)
        # Collapse excessive whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text


def get_parser(path: Path) -> DocumentParser:
    """Get the appropriate parser for a file based on its extension.

    Args:
        path: Path to the document file.

    Returns:
        An instance of the correct DocumentParser subclass.

    Raises:
        ValueError: If no parser supports the file extension.
    """
    parsers: list[DocumentParser] = [
        PDFParser(),
        DOCXParser(),
        TextParser(),
        HTMLParser(),
    ]
    for parser in parsers:
        if parser.can_handle(path):
            return parser

    supported = set()
    for p in parsers:
        supported.update(p.supported_extensions)

    raise ValueError(
        f"No parser available for '{path.suffix}'. "
        f"Supported formats: {', '.join(sorted(supported))}"
    )
